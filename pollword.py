#coding:utf-8
#利用python-jenkins包获取jenkins信息

import jenkins
import os
import sys
import docx
import time
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
reload(sys)
sys.setdefaultencoding( "utf-8" )

f=open(u"C:\\Users\\admin\\Desktop\\巡检\\number.txt",'a+')

data=time.strftime('%Y年%m月%d日%H时%M分',time.localtime())
jenkins_url=""
user_id=""
api_token=""
job_name='halfhour'

server=jenkins.Jenkins(jenkins_url,username=user_id,password=api_token)
jobInfo = server.get_job_info(job_name)
lastBuildNumber = server.get_job_info(job_name)['lastBuild']['number']
info=server.get_build_console_output(job_name,lastBuildNumber)
f.write(str(lastBuildNumber)+'\r\n')
f.close()
file=docx.Document()
#title1=file.add_paragraph(u"保障期间")
title1=file.add_heading(u"保障期间保障")

title1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER #设置居中

title2=file.add_heading(u"每30分钟巡检结果")
title2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

file.add_paragraph(u"。")
file.add_paragraph(u"—巡检")
file.add_paragraph(info)
title3=file.add_paragraph(u"")
title3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
title4=file.add_paragraph(u"巡检时间: %s"%data)
title4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT


file.save(u"C:\\Users\\admin\\Desktop\\巡检\\，每30分钟一次_%s.docx"%data)














